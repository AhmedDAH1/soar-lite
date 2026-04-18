from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from datetime import datetime
import json
from typing import Optional

from app.models import Incident


class ReportService:
    """
    Generate incident reports in PDF and DOCX formats.
    """
    
    @staticmethod
    def generate_pdf_report(incident: Incident) -> BytesIO:
        """
        Generate comprehensive PDF incident report.
        
        Args:
            incident: Incident object with loaded relationships
            
        Returns:
            BytesIO buffer containing PDF data
        """
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=72, leftMargin=72,
                                topMargin=72, bottomMargin=18)
        
        # Container for report elements
        elements = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1e40af'),
            spaceAfter=30,
            alignment=1  # Center
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#1e40af'),
            spaceAfter=12,
            spaceBefore=12
        )
        
        # Title
        severity_colors = {
            'critical': colors.red,
            'high': colors.orange,
            'medium': colors.yellow,
            'low': colors.blue
        }
        severity_color = severity_colors.get(incident.severity.value, colors.grey)
        
        title = Paragraph(
            f"<b>INCIDENT REPORT</b><br/>"
            f"<font color='{severity_color.hexval()}'>[{incident.severity.value.upper()}]</font> "
            f"{incident.title}",
            title_style
        )
        elements.append(title)
        elements.append(Spacer(1, 0.2 * inch))
        
        # Metadata Table
        metadata = [
            ['Incident ID:', f'#{incident.id}'],
            ['Severity:', incident.severity.value.upper()],
            ['Status:', incident.status.value.upper()],
            ['Created:', incident.created_at.strftime('%Y-%m-%d %H:%M:%S UTC')],
            ['Last Updated:', incident.updated_at.strftime('%Y-%m-%d %H:%M:%S UTC') if incident.updated_at else 'N/A'],
        ]
        
        metadata_table = Table(metadata, colWidths=[2*inch, 4*inch])
        metadata_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#e5e7eb')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey)
        ]))
        elements.append(metadata_table)
        elements.append(Spacer(1, 0.3 * inch))
        
        # Executive Summary
        elements.append(Paragraph("<b>EXECUTIVE SUMMARY</b>", heading_style))
        summary_text = incident.description or "No description available."
        elements.append(Paragraph(summary_text, styles['BodyText']))
        elements.append(Spacer(1, 0.2 * inch))
        
        # Alerts Section
        if incident.alerts:
            elements.append(Paragraph("<b>ALERTS</b>", heading_style))
            
            alert_data = [['Source', 'Title', 'Created']]
            for alert in incident.alerts:
                alert_data.append([
                    alert.source.upper(),
                    alert.title[:50] + '...' if len(alert.title) > 50 else alert.title,
                    alert.created_at.strftime('%Y-%m-%d %H:%M')
                ])
            
            alert_table = Table(alert_data, colWidths=[1*inch, 3.5*inch, 1.5*inch])
            alert_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3b82f6')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('GRID', (0, 0), (-1, -1), 1, colors.grey)
            ]))
            elements.append(alert_table)
            elements.append(Spacer(1, 0.2 * inch))
        
        # IOCs Section
        if incident.iocs:
            elements.append(Paragraph("<b>INDICATORS OF COMPROMISE (IOCs)</b>", heading_style))
            
            ioc_data = [['Type', 'Value', 'Malicious', 'Enrichment Summary']]
            for ioc in incident.iocs:
                # Parse enrichment data for summary
                enrichment_summary = "Not enriched"
                if ioc.enrichment_data:
                    try:
                        enrich = json.loads(ioc.enrichment_data)
                        parts = []
                        if 'virustotal' in enrich:
                            vt = enrich['virustotal']
                            parts.append(f"VT: {vt.get('malicious', 0)} detections")
                        if 'abuseipdb' in enrich:
                            abuse = enrich['abuseipdb']
                            parts.append(f"Abuse: {abuse.get('abuse_confidence_score', 0)}%")
                        if 'geolocation' in enrich:
                            geo = enrich['geolocation']
                            parts.append(f"{geo.get('country', 'Unknown')}")
                        enrichment_summary = ' | '.join(parts) if parts else "Enriched"
                    except:
                        enrichment_summary = "Parse error"
                
                ioc_data.append([
                    ioc.type.value.upper(),
                    ioc.value[:40] + '...' if len(ioc.value) > 40 else ioc.value,
                    '✗ YES' if ioc.is_malicious else '✓ No',
                    enrichment_summary[:30] + '...' if len(enrichment_summary) > 30 else enrichment_summary
                ])
            
            ioc_table = Table(ioc_data, colWidths=[0.7*inch, 2*inch, 0.8*inch, 2.5*inch])
            ioc_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3b82f6')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f3f4f6')])
            ]))
            elements.append(ioc_table)
            elements.append(Spacer(1, 0.2 * inch))
        
        # Timeline Section
        if incident.actions:
            elements.append(PageBreak())
            elements.append(Paragraph("<b>INCIDENT TIMELINE</b>", heading_style))
            
            timeline_data = [['Timestamp', 'Action', 'Performed By']]
            for action in incident.actions:
                timeline_data.append([
                    action.created_at.strftime('%m/%d %H:%M'),
                    action.description[:60] + '...' if len(action.description) > 60 else action.description,
                    action.performed_by
                ])
            
            timeline_table = Table(timeline_data, colWidths=[1*inch, 3.5*inch, 1.5*inch])
            timeline_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3b82f6')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f3f4f6')])
            ]))
            elements.append(timeline_table)
        
        # Footer
        elements.append(Spacer(1, 0.5 * inch))
        footer_text = f"<i>Report generated by SOAR-Lite on {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}</i>"
        elements.append(Paragraph(footer_text, styles['Normal']))
        
        # Build PDF
        doc.build(elements)
        buffer.seek(0)
        return buffer
    
    @staticmethod
    def generate_docx_report(incident: Incident) -> BytesIO:
        """
        Generate editable DOCX incident report.
        
        Args:
            incident: Incident object with loaded relationships
            
        Returns:
            BytesIO buffer containing DOCX data
        """
        doc = Document()
        
        # Title
        title = doc.add_heading('INCIDENT REPORT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading(f'[{incident.severity.value.upper()}] {incident.title}', level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Set subtitle color based on severity
        severity_colors = {
            'critical': RGBColor(220, 38, 38),
            'high': RGBColor(249, 115, 22),
            'medium': RGBColor(234, 179, 8),
            'low': RGBColor(59, 130, 246)
        }
        if incident.severity.value in severity_colors:
            for run in subtitle.runs:
                run.font.color.rgb = severity_colors[incident.severity.value]
        
        doc.add_paragraph()  # Spacer
        
        # Metadata Table
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'
        
        metadata = [
            ('Incident ID:', f'#{incident.id}'),
            ('Severity:', incident.severity.value.upper()),
            ('Status:', incident.status.value.upper()),
            ('Created:', incident.created_at.strftime('%Y-%m-%d %H:%M:%S UTC')),
            ('Last Updated:', incident.updated_at.strftime('%Y-%m-%d %H:%M:%S UTC') if incident.updated_at else 'N/A')
        ]
        
        for i, (key, value) in enumerate(metadata):
            table.rows[i].cells[0].text = key
            table.rows[i].cells[1].text = value
            table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(incident.description or "No description available.")
        
        # Alerts
        if incident.alerts:
            doc.add_heading('Alerts', level=1)
            
            alert_table = doc.add_table(rows=1 + len(incident.alerts), cols=3)
            alert_table.style = 'Light Grid Accent 1'
            
            # Header
            header_cells = alert_table.rows[0].cells
            header_cells[0].text = 'Source'
            header_cells[1].text = 'Title'
            header_cells[2].text = 'Created'
            
            for cell in header_cells:
                cell.paragraphs[0].runs[0].font.bold = True
            
            # Data
            for i, alert in enumerate(incident.alerts, start=1):
                row_cells = alert_table.rows[i].cells
                row_cells[0].text = alert.source.upper()
                row_cells[1].text = alert.title
                row_cells[2].text = alert.created_at.strftime('%Y-%m-%d %H:%M')
        
        # IOCs
        if incident.iocs:
            doc.add_heading('Indicators of Compromise (IOCs)', level=1)
            
            ioc_table = doc.add_table(rows=1 + len(incident.iocs), cols=4)
            ioc_table.style = 'Light Grid Accent 1'
            
            # Header
            header_cells = ioc_table.rows[0].cells
            header_cells[0].text = 'Type'
            header_cells[1].text = 'Value'
            header_cells[2].text = 'Malicious'
            header_cells[3].text = 'Enrichment Summary'
            
            for cell in header_cells:
                cell.paragraphs[0].runs[0].font.bold = True
            
            # Data
            for i, ioc in enumerate(incident.iocs, start=1):
                row_cells = ioc_table.rows[i].cells
                row_cells[0].text = ioc.type.value.upper()
                row_cells[1].text = ioc.value
                row_cells[2].text = '✗ YES' if ioc.is_malicious else '✓ No'
                
                # Enrichment summary
                enrichment_summary = "Not enriched"
                if ioc.enrichment_data:
                    try:
                        enrich = json.loads(ioc.enrichment_data)
                        parts = []
                        if 'virustotal' in enrich:
                            parts.append(f"VT: {enrich['virustotal'].get('malicious', 0)} detections")
                        if 'abuseipdb' in enrich:
                            parts.append(f"Abuse: {enrich['abuseipdb'].get('abuse_confidence_score', 0)}%")
                        enrichment_summary = ' | '.join(parts) if parts else "Enriched"
                    except:
                        enrichment_summary = "Parse error"
                
                row_cells[3].text = enrichment_summary
        
        # Timeline
        if incident.actions:
            doc.add_page_break()
            doc.add_heading('Incident Timeline', level=1)
            
            timeline_table = doc.add_table(rows=1 + len(incident.actions), cols=3)
            timeline_table.style = 'Light Grid Accent 1'
            
            # Header
            header_cells = timeline_table.rows[0].cells
            header_cells[0].text = 'Timestamp'
            header_cells[1].text = 'Action'
            header_cells[2].text = 'Performed By'
            
            for cell in header_cells:
                cell.paragraphs[0].runs[0].font.bold = True
            
            # Data
            for i, action in enumerate(incident.actions, start=1):
                row_cells = timeline_table.rows[i].cells
                row_cells[0].text = action.created_at.strftime('%m/%d %H:%M')
                row_cells[1].text = action.description
                row_cells[2].text = action.performed_by
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.add_run(f"Report generated by SOAR-Lite on {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}").italic = True
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer